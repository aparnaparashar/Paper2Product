import io


def export_markdown(project) -> str:
    fr   = project.final_report or {}
    rp   = project.research_profile or {}
    opp  = project.opportunity_scores or {}
    inv  = project.investment_score or {}
    risk = project.risk_profile or {}
    mo   = project.market_opportunities or {}
    rev  = project.revenue_strategy or {}
    mvp  = project.mvp_plan or {}

    lines = [
        f"# {project.title}",
        f"\n**Verdict:** {fr.get('go_no_go','—')} — {fr.get('overall_verdict','')}\n",
        "---\n",
        "## Executive Summary",
        fr.get("executive_summary", ""),
        "\n---\n",
        "## Opportunity Scores",
    ]
    for k, v in (opp.get("dimensions") or {}).items():
        if isinstance(v, dict):
            lines.append(f"- **{k.replace('_',' ').title()}:** {v.get('score','—')}/10 — {v.get('rationale','')}")
    lines.append(f"\n**Final Score: {opp.get('weighted_final_score','—')}/10**\n")

    lines += [
        "---\n## Research",
        f"- **Domain:** {rp.get('domain','')}",
        f"- **Problem:** {rp.get('problem','')}",
        f"- **Novelty:** {rp.get('novelty','')}",
        "\n---\n## Market",
        f"- **Market:** {(mo.get('primary_market') or {}).get('name','')}  "
        f"${(mo.get('primary_market') or {}).get('size_usd_billion','')}B  "
        f"{(mo.get('primary_market') or {}).get('growth_rate_pct','')}% CAGR",
        f"- **Timing:** {mo.get('market_timing','')}",
        "\n---\n## MVP",
        f"- **Timeline:** {mvp.get('time_to_mvp','')}",
        f"- **Cost:** ${(mvp.get('estimated_cost_usd') or {}).get('min','')}–${(mvp.get('estimated_cost_usd') or {}).get('max','')}",
        "\n---\n## Risk",
        f"- **Overall Risk:** {risk.get('overall_risk_level','')}",
    ]
    for r in (risk.get("top_3_risks") or []):
        lines.append(f"- {r}")

    lines += [
        "\n---\n## Investment",
        f"- **Score:** {inv.get('investment_score','')} / 100",
        f"- **Verdict:** {inv.get('investment_verdict','')}",
        f"\n{inv.get('investment_memo_summary','')}",
        "\n---\n## Agent Debate Resolutions",
    ]
    for c in (fr.get("debate_resolution") or []):
        lines += [
            f"\n### {c.get('conflict','')}",
            f"- Product: {c.get('product_argument','')}",
            f"- Risk: {c.get('risk_argument','')}",
            f"- **Ruling:** {c.get('judge_ruling','')}",
        ]

    lines += ["\n---\n## Final Recommendation"]
    for p in (fr.get("final_recommendation_paragraphs") or []):
        lines.append(f"\n{p}")

    return "\n".join(lines)


def export_pdf(project) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        BRAND = colors.HexColor("#4f6ef7")
        DARK  = colors.HexColor("#111827")
        GRAY  = colors.HexColor("#6b7280")

        title_s = ParagraphStyle("T", parent=styles["Title"], textColor=BRAND, fontSize=22, spaceAfter=6)
        h2_s    = ParagraphStyle("H2", parent=styles["Heading2"], textColor=DARK, fontSize=14, spaceBefore=14, spaceAfter=4)
        body_s  = ParagraphStyle("B", parent=styles["Normal"], fontSize=10, leading=15)

        fr  = project.final_report or {}
        opp = project.opportunity_scores or {}
        inv = project.investment_score or {}
        risk= project.risk_profile or {}

        story = [
            Paragraph(project.title, title_s),
            Paragraph(f"Verdict: <b>{fr.get('go_no_go','')}</b> — {fr.get('overall_verdict','')}", body_s),
            Spacer(1, 8),
            HRFlowable(width="100%", color=BRAND),
            Spacer(1, 8),
            Paragraph("Executive Summary", h2_s),
            Paragraph(fr.get("executive_summary", ""), body_s),
            Spacer(1, 8),
            Paragraph("Opportunity Scores", h2_s),
        ]

        dims = opp.get("dimensions", {})
        if dims:
            data = [["Dimension", "Score", "Rationale"]]
            for k, v in dims.items():
                if isinstance(v, dict):
                    data.append([k.replace("_"," ").title(), f"{v.get('score','—')}/10", (v.get("rationale",""))[:90]])
            t = Table(data, colWidths=[4*cm, 2*cm, 10*cm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), BRAND),
                ("TEXTCOLOR", (0,0), (-1,0), colors.white),
                ("FONTSIZE", (0,0), (-1,-1), 9),
                ("GRID", (0,0), (-1,-1), 0.4, colors.lightgrey),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f9fafb")]),
            ]))
            story += [t, Spacer(1, 6)]
        story.append(Paragraph(f"<b>Final Score: {opp.get('weighted_final_score','—')}/10</b>", body_s))
        story += [Spacer(1, 8), Paragraph("Investment", h2_s),
                  Paragraph(f"Score: {inv.get('investment_score','')}/100 · {inv.get('investment_verdict','')}", body_s),
                  Paragraph(inv.get("investment_memo_summary",""), body_s),
                  Spacer(1, 8), Paragraph("Risk", h2_s),
                  Paragraph(f"Overall: {risk.get('overall_risk_level','')}", body_s)]
        for r in (risk.get("top_3_risks") or []):
            story.append(Paragraph(f"• {r}", body_s))
        story += [Spacer(1,8), Paragraph("Final Recommendation", h2_s)]
        for p in (fr.get("final_recommendation_paragraphs") or []):
            story += [Paragraph(p, body_s), Spacer(1, 5)]

        doc.build(story)
        return buf.getvalue()
    except Exception:
        return export_markdown(project).encode("utf-8")


def export_docx(project) -> bytes:
    try:
        import docx
        from docx.shared import Pt
        doc = docx.Document()
        fr  = project.final_report or {}
        opp = project.opportunity_scores or {}
        risk= project.risk_profile or {}

        doc.add_heading(project.title, 0)
        doc.add_paragraph(f"Verdict: {fr.get('go_no_go','')} — {fr.get('overall_verdict','')}")
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(fr.get("executive_summary", ""))
        doc.add_heading("Opportunity Scores", 1)
        dims = opp.get("dimensions", {})
        if dims:
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Dimension", "Score", "Rationale"
            for k, v in dims.items():
                if isinstance(v, dict):
                    row = tbl.add_row().cells
                    row[0].text = k.replace("_"," ").title()
                    row[1].text = f"{v.get('score','—')}/10"
                    row[2].text = (v.get("rationale",""))[:120]
        doc.add_paragraph(f"Final Score: {opp.get('weighted_final_score','—')}/10")
        doc.add_heading("Risk", 1)
        doc.add_paragraph(f"Overall: {risk.get('overall_risk_level','')}")
        for r in (risk.get("top_3_risks") or []):
            doc.add_paragraph(f"• {r}")
        doc.add_heading("Final Recommendation", 1)
        for p in (fr.get("final_recommendation_paragraphs") or []):
            doc.add_paragraph(p)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return export_markdown(project).encode("utf-8")
